from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = (
    ROOT
    / "app"
    / "downloads"
    / "Aurora_Forge_Complete_WWE_2K25_2K26_PC_Modding_Handbook.docx"
)

REQUIRED_HEADINGS = (
    "Begin Here",
    "Install a ready-made character",
    "Build a Complete Character",
    "Move Content from WWE 2K25 to WWE 2K26",
    "Arenas and Show Presentation",
    "WWE 2K25 notes",
    "Troubleshooting",
    "Source notes and bibliography",
)


def main() -> None:
    document = Document(DOCX_PATH)
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    ]

    with ZipFile(DOCX_PATH) as archive:
        damaged_member = archive.testzip()
        media = [
            name for name in archive.namelist() if name.startswith("word/media/")
        ]

    hyperlinks = [
        relationship
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink")
    ]

    if damaged_member is not None:
        raise RuntimeError(f"DOCX ZIP member is damaged: {damaged_member}")

    missing = [
        required
        for required in REQUIRED_HEADINGS
        if not any(required.lower() in heading.lower() for heading in headings)
    ]
    if missing:
        nearby = [
            heading
            for heading in headings
            if any(
                token in heading.lower()
                for token in ("install", "character", "port", "arena", "troubleshooting")
            )
        ]
        print(f"candidate_headings={nearby}")
        raise RuntimeError(f"Required DOCX sections are missing: {missing}")

    if not any(
        "shiny" in heading.lower() or "black" in heading.lower()
        for heading in headings
    ):
        raise RuntimeError("Shiny/black character workaround section is missing")

    print(f"paragraphs={len(document.paragraphs)}")
    print(f"headings={len(headings)}")
    print(f"tables={len(document.tables)}")
    print(f"images={len(media)}")
    print(f"hyperlinks={len(hyperlinks)}")
    print("zip_integrity=PASS")
    print(f"first_heading={headings[0]}")
    print(f"last_heading={headings[-1]}")
    print("required_handbook_sections=PASS")


if __name__ == "__main__":
    main()
