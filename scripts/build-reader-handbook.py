#!/usr/bin/env python3
"""Build the public Aurora Forge handbook in Markdown, HTML, PDF, and DOCX.

The original research source is preserved outside the packaged app. Public
outputs deliberately omit internal product-planning and parser-architecture
notes.
"""

from __future__ import annotations

import html
import os
import re
import shutil
import sys
import textwrap
from pathlib import Path
from urllib.parse import quote

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    LongTable,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


ROOT = Path(__file__).resolve().parents[1]
APP = ROOT / "app"
DOWNLOADS = APP / "downloads"
SOURCE_MD = DOWNLOADS / "Aurora_Forge_Complete_WWE_2K25_2K26_PC_Modding_Handbook.md"
SOURCE_DOCX = DOWNLOADS / "Aurora_Forge_Complete_WWE_2K25_2K26_PC_Modding_Handbook.docx"
PDF_OUT = DOWNLOADS / "Aurora_Forge_Complete_WWE_2K25_2K26_PC_Modding_Handbook.pdf"
HTML_OUT = APP / "handbook-reader.html"
DEV_DIR = ROOT / "dev-notes"
RAW_BACKUP = DEV_DIR / "Aurora_Forge_Knowledgebase_Source_v2.0_with_Development_Notes.md"
REMOVED_NOTES = DEV_DIR / "Aurora_Forge_Knowledgebase_Removed_Development_Sections.md"
COMMUNITY_VIDEO_UPDATE = DEV_DIR / "Community_Tutorial_Source_Update_2026-07-25.md"
ARENA_BUILDER_UPDATE = DEV_DIR / "Arena_Builder_Source_Update_2026-07-26.md"
LOGO = APP / "assets" / "img" / "aurora-forge-mark.png"

TITLE = "Aurora Forge: The Complete WWE 2K25 & WWE 2K26 PC Modding Handbook"
SUBTITLE = "A patient, start-to-finish manual for complete beginners"


def read_original_source() -> str:
    DEV_DIR.mkdir(parents=True, exist_ok=True)
    if RAW_BACKUP.exists():
        return RAW_BACKUP.read_text(encoding="utf-8")
    raw = SOURCE_MD.read_text(encoding="utf-8")
    RAW_BACKUP.write_text(raw, encoding="utf-8", newline="\n")
    return raw


def read_public_arena_builder_update() -> str:
    if not ARENA_BUILDER_UPDATE.exists():
        return ""

    source = ARENA_BUILDER_UPDATE.read_text(encoding="utf-8")
    sources_start = source.find("## Sources")
    opportunities_start = source.find("## Aurora Forge opportunities")
    evidence_start = source.find("## Evidence boundary")

    if min(sources_start, opportunities_start, evidence_start) < 0:
        raise RuntimeError("Arena Builder source update is missing a public-section marker")

    public_body = (
        source[sources_start:opportunities_start].strip()
        + "\n\n"
        + source[evidence_start:].strip()
    )
    public_body = re.sub(r"^## ", "#### ", public_body, flags=re.M)
    public_body = re.sub(r"^### ", "##### ", public_body, flags=re.M)

    return (
        "### Arena Builder tutorial update - July 26, 2026\n\n"
        + public_body.strip()
    )


def remove_numbered_section(text: str, start_number: int, next_number: int | None) -> tuple[str, str]:
    if next_number is None:
        pattern = rf"(?ms)^## {start_number}\.\s.*\Z"
    else:
        pattern = rf"(?ms)^## {start_number}\.\s.*?(?=^## {next_number}\.\s)"
    match = re.search(pattern, text)
    if not match:
        return text, ""
    removed = match.group(0).rstrip()
    return text[: match.start()] + text[match.end() :], removed


def remove_opportunity_blocks(text: str) -> tuple[str, list[str]]:
    removed: list[str] = []
    pattern = re.compile(
        r"(?ms)\n\*\*Aurora Forge opportunities\*\*\s*\n.*?"
        r"(?=\n\*\*[^*\n]+\*\*\s*\n|\n##\s)"
    )

    def repl(match: re.Match[str]) -> str:
        removed.append(match.group(0).strip())
        return "\n"

    return pattern.sub(repl, text), removed


def replace_screen_guide(text: str) -> str:
    start = text.find("### 22.1 Home")
    end = text.find("### 22.5 Face Texture Studio")
    if start < 0 or end < 0 or end <= start:
        return text

    replacement = """### 22.1 Home

**Purpose:** choose the part of Aurora Forge you need.

**Use**

1. Open **Projects** when you want to start, save, or continue a project.
2. Open **Creative Studios** when you want to make a character, face, mask, tattoo, gear set, arena package, music plan, or another creative asset.
3. Open **Tools** when you need the Character Viewer, Face Calibration, or Mod File Inspector.
4. Open **Tutorials** when you want short lessons, videos, troubleshooting, or this complete handbook.
5. Open **Setup** to choose folders, connect external programs, and change app options.
6. Open **About** only when you need the version, credits, licenses, or the private Dev panel.

**Output:** none. Home helps you choose where to go.

### 22.2 Projects

**Purpose:** keep the project folder, saved project file, revisions, exports, screenshots, and next actions together.

**Use**

1. Create a project and save it right away.
2. Put untouched copies in **Originals**.
3. Put editable copies in **Work**.
4. Put checked, finished files in **Ready**.
5. Put in-game pictures in **Screenshots**.
6. Duplicate the project before a risky change.
7. Export a backup after every good in-game test.

**Stop if:** you cannot tell which files are original, you are editing inside the game folder, or you do not have a backup.

### 22.3 Complete Character

**Purpose:** plan a whole character so the face, body, tattoos, gear, music, entrance, victory, and announcer notes all match.

**Use**

1. Fill in the character name and simple identity details.
2. Choose the official CAW path or the file-level PC mod path.
3. Plan the face and hair.
4. Plan tattoos and body details.
5. Plan ring gear and entrance gear.
6. Choose colors, logos, materials, music, entrance, victory, and announcer notes.
7. Save the project and download the final handoff pack.
8. Make each asset in its matching Creative Studio.
9. Install and test one part at a time.

**Output:** one coordinated project and its handoff files. It does not register a finished character in the game by itself.

> **VIDEO NEEDED - AF-V10: Complete Character start to finish, 4-6 minutes**  
> Use one original character. Show the Projects page, Complete Character, each matching studio, the final handoff ZIP, and the safe test-and-fix loop.

"""
    result = text[:start] + replacement + text[end:]
    # Removing the old Workspace subsection shifts the remaining 22.x screen
    # numbers down by one. The top-level section is renumbered later.
    def shift(match: re.Match[str]) -> str:
        number = int(match.group(1))
        if number >= 5:
            return f"### 22.{number - 1} "
        return match.group(0)

    result = re.sub(r"^### 22\.(\d+)\s+", shift, result, flags=re.M)
    return result


def renumber_sections(text: str) -> str:
    mapping = {n: n for n in range(1, 10)}
    mapping.update({n: n - 2 for n in range(12, 31)})

    def heading_repl(match: re.Match[str]) -> str:
        hashes, first, tail, title = match.groups()
        old = int(first)
        new = mapping.get(old, old)
        return f"{hashes} {new}{tail or ''}. {title}"

    text = re.sub(
        r"^(#{2,6})\s+(\d+)((?:\.\d+)*)\.?\s+(.+)$",
        heading_repl,
        text,
        flags=re.M,
    )

    def reference_repl(match: re.Match[str]) -> str:
        word, first, tail = match.groups()
        old = int(first)
        new = mapping.get(old, old)
        return f"{word} {new}{tail or ''}"

    return re.sub(
        r"\b([Ss]ection)\s+(\d+)((?:\.\d+)*)",
        reference_repl,
        text,
    )


def curate_markdown(raw: str) -> str:
    removed_sections: list[str] = []
    text = raw
    supplements: list[str] = []
    if COMMUNITY_VIDEO_UPDATE.exists():
        supplements.append(
            COMMUNITY_VIDEO_UPDATE.read_text(encoding="utf-8").strip()
        )
    arena_builder_supplement = read_public_arena_builder_update()
    if arena_builder_supplement:
        supplements.append(arena_builder_supplement)
    tutorial_policy = "\n### YouTube index policy"
    if supplements and tutorial_policy in text:
        text = text.replace(
            tutorial_policy,
            "\n\n" + "\n\n".join(supplements) + "\n\n### YouTube index policy",
            1,
        )
    text, removed = remove_numbered_section(text, 10, 12)
    if removed:
        removed_sections.append(removed)
    text, removed = remove_numbered_section(text, 31, None)
    if removed:
        removed_sections.append(removed)
    text, opportunity_notes = remove_opportunity_blocks(text)

    notes = [
        "# Aurora Forge internal knowledgebase notes",
        "",
        "These sections were removed from the public Reader Edition because they are product planning or parser architecture notes.",
        "",
    ]
    notes.extend(removed_sections)
    if opportunity_notes:
        notes.extend(["", "## Removed per-tool capability ideas", ""])
        notes.extend(opportunity_notes)
    REMOVED_NOTES.write_text("\n\n".join(notes).strip() + "\n", encoding="utf-8", newline="\n")

    text = text.replace(
        "**Living knowledgebase, start-to-finish tutorial, and Aurora Forge development reference**",
        "**Living knowledgebase and start-to-finish tutorial for Aurora Forge users**",
    )
    text = text.replace(
        "**Edition:** 2.0 — Complete Workflow Edition",
        "**Edition:** 2.2 - Reader Edition for Aurora Forge v1.6.0.d",
    )
    text = text.replace(
        "The safest Aurora Forge strategy is therefore to become the **workflow intelligence and validation layer** around these tools, not to imitate or redistribute them. High-value near-term capabilities are read-only inspection, folder comparison, DDS validation, texture-set completeness checks, profile-driven handoff packs, external-tool path management, and evidence-linked tutorials. Binary editing of MCD, MTLS, JMTL, or YCL should remain out of scope until formats are documented through repeatable tests.",
        "Aurora Forge works as the guide and validation layer around these tools. Use **Projects** to protect originals, **Creative Studios** to plan assets, **Tools** to preview or inspect, **Tutorials** for step-by-step help, and **Setup** for folders and external programs. Keep MCD, MTLS, JMTL, and YCL work inside trusted tools unless a procedure has been tested for the exact game and tool version.",
    )
    text = text.replace(
        "> **Visual-production convention:** Items marked **PHOTO NEEDED** or **VIDEO NEEDED** are production briefs for the user. They identify the exact real-tool screenshot or clip that should later be inserted. They are not missing instructions; they prevent this handbook from using invented interface screenshots.",
        "> **Pictures and videos:** Items marked **PHOTO NEEDED** or **VIDEO NEEDED** tell the tutorial team exactly what real screenshot or clip to add. The written steps still work without them.",
    )

    quick_start = """
> **New here? Start small.** Open **Tutorials** in Aurora Forge, choose one easy lesson, and finish one step at a time. Make a backup before touching game files. Stop when a screen or file does not match the lesson.
"""
    marker = "> **Pictures and videos:**"
    marker_end = text.find("\n\n", text.find(marker))
    if marker_end >= 0:
        text = text[:marker_end] + "\n\n" + quick_start.strip() + text[marker_end:]

    text = replace_screen_guide(text)
    replacements = {
        "Complete CAW Builder": "Complete Character",
        "Create a Complete CAW project.": "Create a Complete Character project.",
        "Tutorials and Training Pack": "Tutorials",
        "Experimental Live Character Viewer": "Character Viewer",
        "Modding Framework": "File Chain Reference",
        "This section defines what every visible Aurora Forge capability should teach, what it produces, and where the user goes next.": (
            "Aurora Forge has seven main sections: Home, Projects, Creative Studios, Tools, Tutorials, Setup, and About. "
            "The steps below explain what each workspace does and what you should have before moving on."
        ),
        "Visual-production uploader": "Visual tutorial library",
        "development reference": "user reference",
        "development-facing": "advanced",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Add Setup and About help immediately before the next major workflow.
    insert_at = text.find("\n## 23. Official CAW")
    if insert_at >= 0:
        extra = """

### 22.21 Setup

**Purpose:** choose Aurora Forge folders, connect external programs, and keep all location settings in one place.

1. Choose the Projects folder.
2. Choose the Exports folder.
3. Choose the WWE 2K26 game folder.
4. Add only the external programs you actually use.
5. Press each check button.
6. Fix every red result before starting a file-level mod.

Aurora Forge remembers paths; it does not bundle paid tools or copy them into the app.

### 22.22 About

**Purpose:** show the app version, credits, licenses, and release information.

The **Dev** panel is intentionally separate. Normal users do not need it for projects, studios, tools, tutorials, or setup.
"""
        text = text[:insert_at] + extra + text[insert_at:]

    text = renumber_sections(text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip() + "\n"


def slugify(value: str, used: dict[str, int]) -> str:
    clean = re.sub(r"<[^>]+>", "", value)
    clean = re.sub(r"[^\w\s-]", "", clean, flags=re.UNICODE).strip().lower()
    clean = re.sub(r"[-\s]+", "-", clean) or "section"
    count = used.get(clean, 0)
    used[clean] = count + 1
    return clean if count == 0 else f"{clean}-{count + 1}"


INLINE_TOKEN = re.compile(
    r"(\[[^\]]+\]\([^)]+\)|`[^`]+`|\*\*[^*]+\*\*|\*[^*\n]+\*)"
)


def render_inline_html(text: str) -> str:
    parts: list[str] = []
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        parts.append(html.escape(text[position : match.start()]))
        token = match.group(0)
        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link:
            label = html.escape(link.group(1))
            href = html.escape(link.group(2), quote=True)
            parts.append(f'<a href="{href}" target="_blank" rel="noreferrer">{label}</a>')
        elif token.startswith("`"):
            parts.append(f"<code>{html.escape(token[1:-1])}</code>")
        elif token.startswith("**"):
            parts.append(f"<strong>{html.escape(token[2:-2])}</strong>")
        else:
            parts.append(f"<em>{html.escape(token[1:-1])}</em>")
        position = match.end()
    parts.append(html.escape(text[position:]))
    return "".join(parts).replace("  \n", "<br>")


def parse_table(lines: list[str], index: int) -> tuple[str, int]:
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in rows[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0] if rows else []
        body = rows[1:]
    output = ['<div class="handbook-table-wrap"><table><thead><tr>']
    output.extend(f"<th>{render_inline_html(cell)}</th>" for cell in header)
    output.append("</tr></thead><tbody>")
    for row in body:
        output.append("<tr>")
        padded = row + [""] * (len(header) - len(row))
        output.extend(f"<td>{render_inline_html(cell)}</td>" for cell in padded[: len(header)])
        output.append("</tr>")
    output.append("</tbody></table></div>")
    return "".join(output), index


def markdown_to_html(markdown: str) -> tuple[str, list[tuple[int, str, str]]]:
    lines = markdown.splitlines()
    output: list[str] = []
    toc: list[tuple[int, str, str]] = []
    used: dict[str, int] = {}
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                output.append("<pre><code>" + html.escape("\n".join(code_lines)) + "</code></pre>")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            raw_level = len(heading.group(1))
            title = heading.group(2).strip()
            numeric = re.match(r"^(\d+(?:\.\d+)*)\.?\s+", title)
            level = min(4, 1 + (numeric.group(1).count(".") if numeric else max(0, raw_level - 1)))
            anchor = slugify(title, used)
            if raw_level > 1 or title.startswith("Part "):
                toc.append((level, title, anchor))
            output.append(f'<h{level} id="{anchor}">{render_inline_html(title)}</h{level}>')
            index += 1
            continue
        if line.lstrip().startswith("|"):
            table_html, index = parse_table(lines, index)
            output.append(table_html)
            continue
        if line.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].startswith(">"):
                quote_lines.append(lines[index][1:].lstrip())
                index += 1
            quote_text = "\n".join(quote_lines)
            kind = "visual" if ("PHOTO NEEDED" in quote_text or "VIDEO NEEDED" in quote_text) else "note"
            output.append(f'<aside class="handbook-callout {kind}">{render_inline_html(quote_text)}</aside>')
            continue
        list_match = re.match(r"^\s*([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            ordered = list_match.group(1)[0].isdigit()
            tag = "ol" if ordered else "ul"
            items: list[str] = []
            while index < len(lines):
                current = re.match(r"^\s*([-*]|\d+\.)\s+(.+)$", lines[index])
                if not current or current.group(1)[0].isdigit() != ordered:
                    break
                items.append(current.group(2))
                index += 1
            output.append(f"<{tag}>" + "".join(f"<li>{render_inline_html(item)}</li>" for item in items) + f"</{tag}>")
            continue
        if re.fullmatch(r"-{3,}", line.strip()):
            output.append("<hr>")
            index += 1
            continue

        paragraph = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if not candidate.strip():
                break
            if re.match(r"^(#{1,6})\s+|^```|^>|^\s*([-*]|\d+\.)\s+|^\s*\|", candidate):
                break
            paragraph.append(candidate.strip())
            index += 1
        output.append("<p>" + render_inline_html(" ".join(paragraph)) + "</p>")

    return "\n".join(output), toc


def build_reader_html(markdown: str) -> None:
    body, toc = markdown_to_html(markdown)
    toc_links = []
    for level, title, anchor in toc:
        toc_links.append(
            f'<a class="handbook-toc-level-{min(level, 3)}" href="#{anchor}">{html.escape(title)}</a>'
        )
    page = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<meta http-equiv="Content-Security-Policy" content="default-src 'self'; script-src 'self'; style-src 'self' 'unsafe-inline'; img-src 'self' data: blob:; media-src 'self' blob:; connect-src 'self'; font-src 'self' data:; object-src 'none'; base-uri 'self'; form-action 'none'">
<title>Aurora Forge - Complete Handbook</title>
<link rel="stylesheet" href="assets/css/style.css?v=20260725-r160c">
</head>
<body class="desktop-studio-shell handbook-reader-page">
<div class="app-shell">
  <aside class="app-sidebar">
    <div class="app-brand"><div class="app-mark app-mark-logo"><img src="assets/img/aurora-forge-mark.png" alt="Aurora Forge logo"></div><div><div class="app-brand-title">Aurora Forge</div><div class="app-brand-sub">by VikingStudios</div></div></div>
    <nav class="app-side-nav" aria-label="Primary app navigation">
      <a class="app-nav-link" href="index.html"><span class="nav-ico">⌂</span><span><strong>Home</strong><small>Start here</small></span></a>
      <a class="app-nav-link" href="project-manager.html"><span class="nav-ico">▣</span><span><strong>Projects</strong><small>Save and continue</small></span></a>
      <a class="app-nav-link" href="creative-studios.html"><span class="nav-ico">◆</span><span><strong>Creative Studios</strong><small>Make assets</small></span></a>
      <a class="app-nav-link" href="tools.html"><span class="nav-ico">⚙</span><span><strong>Tools</strong><small>Preview and inspect</small></span></a>
      <a class="app-nav-link active" href="tutorials.html"><span class="nav-ico">▤</span><span><strong>Tutorials</strong><small>Learn step by step</small></span></a>
      <a class="app-nav-link" href="setup.html"><span class="nav-ico">☷</span><span><strong>Setup</strong><small>Folders and programs</small></span></a>
      <div class="app-nav-divider" aria-hidden="true"></div>
      <a class="app-nav-link" href="about.html"><span class="nav-ico">ⓘ</span><span><strong>About</strong><small>Version and credits</small></span></a>
    </nav>
    <div class="app-sidebar-card"><div class="mini-label">Reading tip</div><strong>Search one word</strong><span>Try CakeView, face, mask, music, arena, MCD, or error.</span></div>
  </aside>
  <main class="app-main handbook-main">
    <section class="app-topbar">
      <div><div class="eyebrow">Aurora Forge · Complete Reader</div><h1>WWE 2K25 &amp; WWE 2K26 PC Modding Handbook</h1><p class="sub">Direct, beginner-friendly instructions from first setup to advanced projects.</p></div>
      <div class="topbar-badges"><span class="version-badge">Handbook 3.0</span><span class="app-runtime-badge">Beginner Edition</span></div>
    </section>
    <section class="handbook-reader-toolbar" aria-label="Handbook controls">
      <a class="ai-btn secondary" href="tutorials.html">Back to Tutorials</a>
      <a class="ai-btn primary" href="downloads/Aurora_Forge_Complete_WWE_2K25_2K26_PC_Modding_Handbook.pdf" download>Download PDF</a>
      <button class="ai-btn secondary" id="handbook-font-smaller" type="button">Smaller text</button>
      <button class="ai-btn secondary" id="handbook-font-larger" type="button">Larger text</button>
      <button class="ai-btn secondary" id="handbook-back-top" type="button">Back to top</button>
    </section>
    <div class="handbook-layout">
      <aside class="handbook-toc-panel">
        <label for="handbook-search">Search this handbook</label>
        <input id="handbook-search" type="search" placeholder="Type a word or file name" autocomplete="off">
        <div class="handbook-search-status" id="handbook-search-status" aria-live="polite">All sections shown</div>
        <nav class="handbook-toc" aria-label="Handbook contents">
          {''.join(toc_links)}
        </nav>
      </aside>
      <article class="handbook-article" id="handbook-article">
        {body}
      </article>
    </div>
    <footer class="app-footer">Aurora Forge · VikingStudios · Complete Handbook</footer>
  </main>
</div>
<script src="assets/js/app-config-loader.js?v=20260725-r160c"></script>
<script src="assets/js/desktop-bridge.js?v=20260725-r160c"></script>
<script src="assets/js/handbook-reader.js?v=20260725-r160c"></script>
</body>
</html>
"""
    HTML_OUT.write_text(page, encoding="utf-8", newline="\n")


def ascii_pdf_text(value: str) -> str:
    replacements = {
        "\u2014": " - ",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2192": "->",
        "\u2193": "down",
        "\u00d7": "x",
        "\u2026": "...",
        "\u00b7": " / ",
        "\u2713": "[PASS]",
        "\u2705": "[PASS]",
        "\u274c": "[STOP]",
        "\u26a0": "Warning:",
        "\u00a0": " ",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def inline_reportlab(text: str) -> str:
    text = ascii_pdf_text(text)
    parts: list[str] = []
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        parts.append(html.escape(text[position : match.start()]))
        token = match.group(0)
        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link:
            label = html.escape(ascii_pdf_text(link.group(1)))
            href = html.escape(link.group(2), quote=True)
            parts.append(f'<link href="{href}" color="#1B6CA8"><u>{label}</u></link>')
        elif token.startswith("`"):
            parts.append(f'<font name="AFMono" color="#17324D">{html.escape(token[1:-1])}</font>')
        elif token.startswith("**"):
            parts.append(f"<b>{html.escape(token[2:-2])}</b>")
        else:
            parts.append(f"<i>{html.escape(token[1:-1])}</i>")
        position = match.end()
    parts.append(html.escape(text[position:]))
    return "".join(parts)


class HandbookDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, **kwargs):
        super().__init__(filename, **kwargs)
        self._bookmark_index = 0
        self._outline_level = -1

    def beforeDocument(self):
        # multiBuild makes more than one layout pass. Stable bookmark keys are
        # required so the generated table of contents can converge.
        self._bookmark_index = 0
        self._outline_level = -1
        return super().beforeDocument()

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and flowable.style.name.startswith("Heading"):
            level = int(flowable.style.name[-1])
            text = flowable.getPlainText()
            key = f"heading-{self._bookmark_index}"
            self._bookmark_index += 1
            self.canv.bookmarkPage(key)
            desired_outline_level = level - 1
            # Some community-authored source sections skip a visible heading
            # depth. PDF outlines cannot jump over a parent level, so clamp the
            # jump while keeping the displayed heading and TOC depth intact.
            outline_level = min(desired_outline_level, self._outline_level + 1)
            outline_level = max(0, outline_level)
            self.canv.addOutlineEntry(text, key, level=outline_level, closed=outline_level > 0)
            self._outline_level = outline_level
            self.notify("TOCEntry", (level, text, self.page, key))


def register_fonts() -> None:
    font_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    fonts = {
        "AFSans": font_dir / "segoeui.ttf",
        "AFSans-Bold": font_dir / "segoeuib.ttf",
        "AFSans-Italic": font_dir / "segoeuii.ttf",
        "AFMono": font_dir / "consola.ttf",
    }
    for name, path in fonts.items():
        if path.exists():
            pdfmetrics.registerFont(TTFont(name, str(path)))
    pdfmetrics.registerFontFamily(
        "AFSans",
        normal="AFSans",
        bold="AFSans-Bold",
        italic="AFSans-Italic",
        boldItalic="AFSans-Bold",
    )


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "Body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="AFSans",
            fontSize=9.4,
            leading=13.0,
            textColor=colors.HexColor("#1D2939"),
            spaceAfter=7,
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=True,
        ),
        "List": ParagraphStyle(
            "List",
            parent=base["BodyText"],
            fontName="AFSans",
            fontSize=9.4,
            leading=12.4,
            textColor=colors.HexColor("#1D2939"),
            spaceAfter=2,
            allowWidows=0,
            allowOrphans=0,
            splitLongWords=True,
        ),
        "Heading1": ParagraphStyle(
            "Heading1",
            fontName="AFSans-Bold",
            fontSize=17,
            leading=21,
            textColor=colors.HexColor("#102A43"),
            spaceBefore=15,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "Heading2": ParagraphStyle(
            "Heading2",
            fontName="AFSans-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#006B8F"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "Heading3": ParagraphStyle(
            "Heading3",
            fontName="AFSans-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#7030A0"),
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "Heading4": ParagraphStyle(
            "Heading4",
            fontName="AFSans-Bold",
            fontSize=9.7,
            leading=12,
            textColor=colors.HexColor("#334E68"),
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "Callout": ParagraphStyle(
            "Callout",
            fontName="AFSans",
            fontSize=9.1,
            leading=12.5,
            textColor=colors.HexColor("#102A43"),
            backColor=colors.HexColor("#EAF7FB"),
            borderColor=colors.HexColor("#44B6D1"),
            borderWidth=0.8,
            borderPadding=8,
            spaceBefore=5,
            spaceAfter=9,
        ),
        "VisualCallout": ParagraphStyle(
            "VisualCallout",
            fontName="AFSans",
            fontSize=8.9,
            leading=12.2,
            textColor=colors.HexColor("#4A2A68"),
            backColor=colors.HexColor("#F4EDFA"),
            borderColor=colors.HexColor("#A56CC1"),
            borderWidth=0.8,
            borderPadding=8,
            # Keep visual-production notes clearly separated from the last
            # numbered tutorial step at page boundaries.
            spaceBefore=10,
            spaceAfter=9,
        ),
        "Code": ParagraphStyle(
            "Code",
            fontName="AFMono",
            fontSize=7.4,
            leading=10.2,
            # Preformatted blocks do not consistently paint their ParagraphStyle
            # background in every ReportLab build. Keep the text dark enough to
            # remain readable even when the light background is omitted.
            textColor=colors.HexColor("#243B53"),
            backColor=colors.HexColor("#EEF4F8"),
            borderPadding=7,
            spaceBefore=4,
            spaceAfter=8,
        ),
        "Small": ParagraphStyle(
            "Small",
            fontName="AFSans",
            fontSize=7.2,
            leading=9.2,
            textColor=colors.HexColor("#52606D"),
        ),
        "TOCHeading": ParagraphStyle(
            "TOCHeading",
            fontName="AFSans-Bold",
            fontSize=24,
            leading=28,
            textColor=colors.HexColor("#102A43"),
            spaceAfter=16,
        ),
    }


def page_header_footer(canvas, doc):
    canvas.saveState()
    width, height = LETTER
    canvas.setStrokeColor(colors.HexColor("#D5E4EC"))
    canvas.setLineWidth(0.5)
    canvas.line(doc.leftMargin, height - 0.55 * inch, width - doc.rightMargin, height - 0.55 * inch)
    canvas.setFont("AFSans", 7.5)
    canvas.setFillColor(colors.HexColor("#52606D"))
    canvas.drawString(doc.leftMargin, height - 0.42 * inch, "AURORA FORGE  /  WWE 2K25 + 2K26 PC MODDING HANDBOOK")
    canvas.drawRightString(width - doc.rightMargin, 0.42 * inch, f"PAGE {doc.page}")
    canvas.restoreState()


def cover_page(canvas, doc):
    canvas.saveState()
    width, height = LETTER
    canvas.setFillColor(colors.HexColor("#07131F"))
    canvas.rect(0, 0, width, height, stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor("#0A2132"))
    canvas.circle(width * 0.86, height * 0.78, 180, stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor("#241638"))
    canvas.circle(width * 0.12, height * 0.16, 145, stroke=0, fill=1)
    canvas.setStrokeColor(colors.HexColor("#1EC8E5"))
    canvas.setLineWidth(2)
    canvas.line(0.9 * inch, 8.15 * inch, 3.15 * inch, 8.15 * inch)
    if LOGO.exists():
        canvas.drawImage(str(LOGO), 0.9 * inch, 8.35 * inch, width=0.72 * inch, height=0.72 * inch, mask="auto", preserveAspectRatio=True)
    canvas.setFillColor(colors.HexColor("#8EEBFA"))
    canvas.setFont("AFSans-Bold", 11)
    canvas.drawString(0.9 * inch, 7.72 * inch, "AURORA FORGE  /  VIKINGSTUDIOS")
    canvas.setFillColor(colors.white)
    canvas.setFont("AFSans-Bold", 28)
    title_lines = ["COMPLETE WWE 2K25 + 2K26", "PC MODDING HANDBOOK"]
    y = 6.65 * inch
    for line in title_lines:
        canvas.drawString(0.9 * inch, y, line)
        y -= 0.46 * inch
    canvas.setFillColor(colors.HexColor("#CBD5E1"))
    canvas.setFont("AFSans", 13)
    canvas.drawString(0.92 * inch, 5.45 * inch, "A patient, start-to-finish manual for complete beginners")
    canvas.setFont("AFSans", 10)
    canvas.drawString(0.92 * inch, 4.98 * inch, "Tools  /  Files  /  Characters  /  Textures  /  Arenas  /  Troubleshooting")
    canvas.setFillColor(colors.HexColor("#1EC8E5"))
    canvas.roundRect(0.9 * inch, 2.0 * inch, 2.05 * inch, 0.38 * inch, 0.08 * inch, stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor("#07131F"))
    canvas.setFont("AFSans-Bold", 9)
    canvas.drawCentredString(1.925 * inch, 2.135 * inch, "HANDBOOK EDITION 3.0")
    canvas.setFillColor(colors.HexColor("#A7B7C8"))
    canvas.setFont("AFSans", 8.5)
    canvas.drawString(0.92 * inch, 1.48 * inch, "Research checked through July 27, 2026")
    canvas.drawString(0.92 * inch, 1.22 * inch, "Independent community handbook - not an official 2K or tool-vendor manual")
    canvas.restoreState()


def numeric_heading_level(title: str, markdown_level: int) -> int:
    match = re.match(r"^(\d+(?:\.\d+)*)\.?\s+", title)
    if match:
        return min(4, 1 + match.group(1).count("."))
    if title.startswith("Part "):
        return 1
    return min(4, max(1, markdown_level))


def wrap_code(code: str, width: int = 92) -> str:
    output: list[str] = []
    for line in ascii_pdf_text(code).splitlines():
        if len(line) <= width:
            output.append(line)
        else:
            indent = re.match(r"^\s*", line).group(0)
            output.extend(
                textwrap.wrap(
                    line,
                    width=width,
                    subsequent_indent=indent + "  ",
                    break_long_words=True,
                    break_on_hyphens=False,
                )
            )
    return "\n".join(output)


def markdown_to_pdf_story(markdown: str, styles: dict[str, ParagraphStyle]):
    lines = markdown.splitlines()
    story = []
    index = 0
    in_code = False
    code_lines: list[str] = []
    first_title_skipped = False

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted(wrap_code("\n".join(code_lines)), styles["Code"]))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            md_level = len(heading.group(1))
            title = heading.group(2).strip()
            if md_level == 1 and not first_title_skipped:
                first_title_skipped = True
                index += 1
                continue
            level = numeric_heading_level(title, md_level)
            story.append(Paragraph(inline_reportlab(title), styles[f"Heading{level}"]))
            index += 1
            continue
        if line.lstrip().startswith("|"):
            raw_rows: list[list[str]] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                raw_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            if len(raw_rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in raw_rows[1]):
                header_row = raw_rows[0]
                body_rows = raw_rows[2:]
            else:
                header_row = raw_rows[0]
                body_rows = raw_rows[1:]
            column_count = max(1, len(header_row))
            table_style = ParagraphStyle(
                "TableCell",
                parent=styles["Small"],
                fontSize=6.6 if column_count >= 5 else 7.4,
                leading=8.2 if column_count >= 5 else 9.2,
                textColor=colors.HexColor("#243B53"),
            )
            header_style = ParagraphStyle(
                "TableHeader",
                parent=table_style,
                fontName="AFSans-Bold",
                textColor=colors.HexColor("#102A43"),
            )
            data = [[Paragraph(inline_reportlab(cell), header_style) for cell in header_row]]
            for row in body_rows:
                padded = row + [""] * (column_count - len(row))
                data.append([Paragraph(inline_reportlab(cell), table_style) for cell in padded[:column_count]])
            available = LETTER[0] - 1.3 * inch
            col_widths = [available / column_count] * column_count
            table = LongTable(data, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCEEF5")),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B7C9D6")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAFC")]),
                    ]
                )
            )
            story.extend([Spacer(1, 4), table, Spacer(1, 8)])
            continue
        if line.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].startswith(">"):
                quote_lines.append(lines[index][1:].lstrip())
                index += 1
            quote_text = " ".join(quote_lines)
            style = styles["VisualCallout"] if ("PHOTO NEEDED" in quote_text or "VIDEO NEEDED" in quote_text) else styles["Callout"]
            story.append(Paragraph(inline_reportlab(quote_text), style))
            continue
        list_match = re.match(r"^\s*([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            ordered = list_match.group(1)[0].isdigit()
            items = []
            while index < len(lines):
                current = re.match(r"^\s*([-*]|\d+\.)\s+(.+)$", lines[index])
                if not current or current.group(1)[0].isdigit() != ordered:
                    break
                item_para = Paragraph(inline_reportlab(current.group(2)), styles["List"])
                items.append(ListItem(item_para, leftIndent=10))
                index += 1
            list_options = {
                "bulletType": "1" if ordered else "bullet",
                "leftIndent": 18,
                "bulletFontName": "AFSans",
                "bulletFontSize": 8.5,
                "bulletColor": colors.HexColor("#006B8F"),
                "spaceAfter": 5,
            }
            if ordered:
                list_options["start"] = "1"
            story.append(ListFlowable(items, **list_options))
            continue
        if re.fullmatch(r"-{3,}", line.strip()):
            story.append(Spacer(1, 6))
            index += 1
            continue
        paragraph = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if not candidate.strip():
                break
            if re.match(r"^(#{1,6})\s+|^```|^>|^\s*([-*]|\d+\.)\s+|^\s*\|", candidate):
                break
            paragraph.append(candidate.strip())
            index += 1
        story.append(Paragraph(inline_reportlab(" ".join(paragraph)), styles["Body"]))
    return story


def build_pdf(markdown: str) -> None:
    register_fonts()
    styles = pdf_styles()
    width, height = LETTER
    doc = HandbookDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.62 * inch,
        title=TITLE,
        author="VikingStudios / Aurora Forge",
        subject="WWE 2K25 and WWE 2K26 PC modding handbook and tutorials",
        creator="Aurora Forge",
    )
    body_frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        width - doc.leftMargin - doc.rightMargin,
        height - doc.topMargin - doc.bottomMargin,
        id="body",
    )
    cover_frame = Frame(0, 0, width, height, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0, id="cover")
    doc.addPageTemplates(
        [
            PageTemplate(id="Cover", frames=[cover_frame], onPage=cover_page),
            PageTemplate(id="Body", frames=[body_frame], onPage=page_header_footer),
        ]
    )
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOC1", fontName="AFSans-Bold", fontSize=9.5, leading=13, leftIndent=0, firstLineIndent=0, textColor=colors.HexColor("#102A43"), spaceBefore=3),
        ParagraphStyle("TOC2", fontName="AFSans", fontSize=8.2, leading=11, leftIndent=15, firstLineIndent=0, textColor=colors.HexColor("#334E68")),
        ParagraphStyle("TOC3", fontName="AFSans", fontSize=7.6, leading=10, leftIndent=30, firstLineIndent=0, textColor=colors.HexColor("#52606D")),
        ParagraphStyle("TOC4", fontName="AFSans", fontSize=7.2, leading=9.5, leftIndent=42, firstLineIndent=0, textColor=colors.HexColor("#627D98")),
    ]
    story = [
        Spacer(1, 10.2 * inch),
        NextPageTemplate("Body"),
        PageBreak(),
        Paragraph("Contents", styles["TOCHeading"]),
        Paragraph(
            "Use the PDF bookmarks or click an entry below. Search in your PDF reader with Ctrl+F.",
            styles["Body"],
        ),
        toc,
        PageBreak(),
    ]
    story.extend(markdown_to_pdf_story(markdown, styles))
    doc.multiBuild(story)

    reader = PdfReader(str(PDF_OUT))
    if len(reader.pages) < 20:
        raise RuntimeError(f"Reader PDF unexpectedly short: {len(reader.pages)} pages")
    if not reader.metadata or TITLE not in str(reader.metadata.title):
        raise RuntimeError("Reader PDF metadata is missing the expected title")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.name = "Segoe UI"
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def add_hyperlink(paragraph, label: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1B6CA8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([props, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        link = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
        if link:
            add_hyperlink(paragraph, link.group(1), link.group(2))
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(23, 50, 77)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def configure_docx_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(29, 41, 57)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
        ("Heading 4", 11, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_docx(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_docx_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "AURORA FORGE  /  WWE 2K25 + 2K26 PC MODDING HANDBOOK"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(82, 96, 109)
    add_page_number(section.footer.paragraphs[0])

    # Editorial cover - a manual's cover, not a normal body page.
    for _ in range(5):
        doc.add_paragraph()
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(0.85))
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("AURORA FORGE  /  VIKINGSTUDIOS")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 107, 143)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Complete WWE 2K25 & WWE 2K26\nPC Modding Handbook")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(16, 42, 67)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("A patient, start-to-finish manual for complete beginners")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(112, 48, 160)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Handbook Edition 3.0\nResearch checked through July 27, 2026\nIndependent community handbook").italic = True
    doc.add_page_break()

    toc_title = doc.add_paragraph("Contents", style="Heading 1")
    toc_title.paragraph_format.space_before = Pt(0)
    doc.add_paragraph("Use Word's Navigation pane or search with Ctrl+F. The in-app reader and PDF also include a clickable contents list.")
    headings: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if not heading or len(heading.group(1)) == 1:
            continue
        level = numeric_heading_level(heading.group(2), len(heading.group(1)))
        if level <= 2:
            headings.append((level, heading.group(2)))
    for level, label in headings:
        p = doc.add_paragraph(style="List Bullet" if level == 1 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(label)
    doc.add_page_break()

    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_title = False
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.18)
                p.paragraph_format.right_indent = Inches(0.18)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run(ascii_pdf_text("\n".join(code_lines)))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(23, 50, 77)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            if len(heading.group(1)) == 1 and not skipped_title:
                skipped_title = True
                index += 1
                continue
            level = numeric_heading_level(heading.group(2), len(heading.group(1)))
            doc.add_paragraph(heading.group(2), style=f"Heading {level}")
            index += 1
            continue
        if line.lstrip().startswith("|"):
            rows: list[list[str]] = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in rows[1]):
                header_row, body_rows = rows[0], rows[2:]
            else:
                header_row, body_rows = rows[0], rows[1:]
            columns = max(1, len(header_row))
            table = doc.add_table(rows=1, cols=columns)
            table.style = "Table Grid"
            for col, value in enumerate(header_row):
                cell = table.rows[0].cells[col]
                cell.text = value
                set_cell_shading(cell, "DCEEF5")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)
            for raw_row in body_rows:
                padded = raw_row + [""] * (columns - len(raw_row))
                cells = table.add_row().cells
                for col, value in enumerate(padded[:columns]):
                    cells[col].text = value
                    for run in cells[col].paragraphs[0].runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(7.6 if columns >= 5 else 8.3)
            usable_dxa = 9360
            widths = [usable_dxa // columns] * columns
            widths[-1] += usable_dxa - sum(widths)
            set_table_geometry(table, widths)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].startswith(">"):
                quote_lines.append(lines[index][1:].lstrip())
                index += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4EDFA" if any("NEEDED" in value for value in quote_lines) else "EAF7FB")
            p._p.get_or_add_pPr().append(shading)
            add_inline_runs(p, " ".join(quote_lines))
            continue
        list_match = re.match(r"^\s*([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            ordered = list_match.group(1)[0].isdigit()
            while index < len(lines):
                current = re.match(r"^\s*([-*]|\d+\.)\s+(.+)$", lines[index])
                if not current or current.group(1)[0].isdigit() != ordered:
                    break
                p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
                add_inline_runs(p, current.group(2))
                index += 1
            continue
        if re.fullmatch(r"-{3,}", line.strip()):
            index += 1
            continue
        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            if not candidate.strip():
                break
            if re.match(r"^(#{1,6})\s+|^```|^>|^\s*([-*]|\d+\.)\s+|^\s*\|", candidate):
                break
            paragraph_lines.append(candidate.strip())
            index += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(paragraph_lines))

    props = doc.core_properties
    props.title = TITLE
    props.subject = "WWE 2K25 and WWE 2K26 PC modding handbook and tutorials"
    props.author = "VikingStudios / Aurora Forge"
    props.comments = "Aurora Forge Handbook Edition 3.0"
    doc.save(SOURCE_DOCX)


def verify_outputs(markdown: str) -> None:
    required = [SOURCE_MD, SOURCE_DOCX, PDF_OUT, HTML_OUT]
    for path in required:
        if not path.exists() or path.stat().st_size < 1000:
            raise RuntimeError(f"Missing or unexpectedly small output: {path}")
    forbidden = [
        "PHOTO NEEDED",
        "VIDEO NEEDED",
        "Visual production plan",
        "Watch only the part you need",
        "reviewed July",
        "evidence boundary",
        "AF-V",
        "AF-P",
        "I reviewed",
        "we reviewed",
        "source says",
        "Community-demonstrated",
    ]
    for phrase in forbidden:
        if phrase in markdown:
            raise RuntimeError(f"Public handbook still contains internal or old wording: {phrase}")
    required_sections = [
        "WWE 2K25 & WWE 2K26 PC Modding Handbook",
        "## 10. Install a ready-made character",
        "## 19. Plan the character in Aurora Forge",
        "## 30. What a port really does",
        "## 39. Use Arena Builder in WWE 2K26",
        "## 48. WWE 2K25 notes",
        "## 80. Source notes and bibliography",
    ]
    for phrase in required_sections:
        if phrase not in markdown:
            raise RuntimeError(f"Public handbook is missing required material: {phrase}")
    Document(str(SOURCE_DOCX))
    reader = PdfReader(str(PDF_OUT))
    text_sample = "\n".join((page.extract_text() or "") for page in reader.pages[:6])
    if "WWE 2K25" not in text_sample or "WWE 2K26" not in text_sample or "Begin Here" not in text_sample:
        raise RuntimeError("PDF text extraction did not find expected handbook content")


def main() -> int:
    curated = SOURCE_MD.read_text(encoding="utf-8")
    build_reader_html(curated)
    build_pdf(curated)
    build_docx(curated)
    verify_outputs(curated)
    reader = PdfReader(str(PDF_OUT))
    print(f"Reader handbook built: {len(reader.pages)} PDF pages")
    print(f"Markdown: {SOURCE_MD.stat().st_size:,} bytes")
    print(f"HTML: {HTML_OUT.stat().st_size:,} bytes")
    print(f"PDF: {PDF_OUT.stat().st_size:,} bytes")
    print(f"DOCX: {SOURCE_DOCX.stat().st_size:,} bytes")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"Reader handbook build failed: {exc}", file=sys.stderr)
        raise
